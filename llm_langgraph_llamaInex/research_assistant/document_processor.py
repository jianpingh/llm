"""
文档处理模块
支持多种格式的文档读取和预处理
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter

# 可选依赖的导入
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，支持多种格式的文档读取和处理"""
    
    def __init__(self, chunk_size: int = 1024, chunk_overlap: int = 200, supported_types: List[str] = None):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 文档分块大小
            chunk_overlap: 分块重叠大小
            supported_types: 支持的文件类型列表，如 ['pdf', 'txt', 'md']
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # 设置支持的文件格式
        default_formats = {'.pdf', '.txt', '.md'}
        if DOCX_AVAILABLE:
            default_formats.add('.docx')
        if BS4_AVAILABLE:
            default_formats.add('.html')
        
        if supported_types:
            # 确保格式以点开头
            requested_formats = {f'.{ext}' if not ext.startswith('.') else ext for ext in supported_types}
            # 只支持实际可用的格式
            self.supported_formats = requested_formats.intersection(default_formats)
            
            # 如果请求的格式不被支持，发出警告
            unsupported = requested_formats - default_formats
            if unsupported:
                logger.warning(f"以下格式不被支持或缺少依赖库: {unsupported}")
        else:
            self.supported_formats = default_formats
        
        logger.info(f"支持的文件格式: {self.supported_formats}")
    
    def load_documents(self, data_dir: str) -> List[Document]:
        """
        从指定目录加载所有支持的文档
        
        Args:
            data_dir: 文档目录路径
            
        Returns:
            处理后的文档列表
        """
        documents = []
        data_path = Path(data_dir)
        
        if not data_path.exists():
            logger.warning(f"数据目录不存在: {data_dir}")
            return documents
        
        for file_path in data_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    doc = self._process_file(file_path)
                    if doc:
                        documents.extend(doc)
                        logger.info(f"成功处理文档: {file_path}")
                except Exception as e:
                    logger.error(f"处理文档失败 {file_path}: {str(e)}")
        
        logger.info(f"总共加载了 {len(documents)} 个文档块")
        return documents
    
    def _process_file(self, file_path: Path) -> Optional[List[Document]]:
        """
        处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档对象列表
        """
        file_suffix = file_path.suffix.lower()
        
        if file_suffix == '.pdf':
            return self._process_pdf(file_path)
        elif file_suffix in ['.txt', '.md']:
            return self._process_text(file_path)
        elif file_suffix == '.docx':
            return self._process_docx(file_path)
        elif file_suffix == '.html':
            return self._process_html(file_path)
        else:
            logger.warning(f"不支持的文件格式: {file_suffix}")
            return None
    
    def _process_pdf(self, file_path: Path) -> List[Document]:
        """处理PDF文件"""
        documents = []
        
        try:
            pdf_document = fitz.open(file_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                
                if text.strip():  # 只处理非空页面
                    doc = Document(
                        text=text,
                        metadata={
                            "file_name": file_path.name,
                            "file_path": str(file_path),
                            "page_number": page_num + 1,
                            "document_type": "pdf"
                        }
                    )
                    documents.append(doc)
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"处理PDF文件失败 {file_path}: {str(e)}")
            
        return documents
    
    def _process_text(self, file_path: Path) -> List[Document]:
        """处理文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document(
                text=content,
                metadata={
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "document_type": "text"
                }
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"处理文本文件失败 {file_path}: {str(e)}")
            return []
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        将文档分割成更小的块
        
        Args:
            documents: 原始文档列表
            
        Returns:
            分割后的文档块列表
        """
        nodes = self.text_splitter.get_nodes_from_documents(documents)
        
        # 转换nodes为Documents并保留元数据
        split_docs = []
        for i, node in enumerate(nodes):
            doc = Document(
                text=node.text,
                metadata={
                    **node.metadata,
                    "chunk_id": i,
                    "chunk_size": len(node.text)
                }
            )
            split_docs.append(doc)
        
        logger.info(f"文档分割完成，共 {len(split_docs)} 个文档块")
        return split_docs
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        获取文档统计信息
        
        Args:
            documents: 文档列表
            
        Returns:
            统计信息字典
        """
        if not documents:
            return {"total_documents": 0}
        
        total_chars = sum(len(doc.text) for doc in documents)
        file_types = {}
        file_names = set()
        
        for doc in documents:
            doc_type = doc.metadata.get("document_type", "unknown")
            file_types[doc_type] = file_types.get(doc_type, 0) + 1
            file_names.add(doc.metadata.get("file_name", "unknown"))
        
        return {
            "total_documents": len(documents),
            "total_characters": total_chars,
            "average_chunk_size": total_chars // len(documents),
            "file_types": file_types,
            "unique_files": len(file_names),
            "file_names": list(file_names)
        }
    
    def _process_docx(self, file_path: Path) -> List[Document]:
        """处理 DOCX 文件"""
        if not DOCX_AVAILABLE:
            logger.error("缺少 python-docx 库，无法处理 DOCX 文件")
            return []
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            content = '\n'.join(text_content)
            
            if content.strip():
                document = Document(
                    text=content,
                    metadata={
                        "file_name": file_path.name,
                        "file_path": str(file_path),
                        "document_type": "docx"
                    }
                )
                return [document]
            else:
                logger.warning(f"DOCX 文件内容为空: {file_path}")
                return []
                
        except Exception as e:
            logger.error(f"处理 DOCX 文件失败 {file_path}: {str(e)}")
            return []
    
    def _process_html(self, file_path: Path) -> List[Document]:
        """处理 HTML 文件"""
        if not BS4_AVAILABLE:
            logger.error("缺少 beautifulsoup4 库，无法处理 HTML 文件")
            return []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本内容
            text = soup.get_text()
            
            # 清理文本
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if text.strip():
                document = Document(
                    text=text,
                    metadata={
                        "file_name": file_path.name,
                        "file_path": str(file_path),
                        "document_type": "html"
                    }
                )
                return [document]
            else:
                logger.warning(f"HTML 文件内容为空: {file_path}")
                return []
                
        except Exception as e:
            logger.error(f"处理 HTML 文件失败 {file_path}: {str(e)}")
            return []
